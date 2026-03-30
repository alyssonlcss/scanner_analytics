"""
DOCX document builder for creating Word reports.

This module provides a fluent interface for building Word documents
following ABNT formatting standards.
"""

from typing import List, Tuple, Optional
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import logging

logger = logging.getLogger(__name__)


class DocxBuilder:
    """
    Builder class for creating Word documents.
    
    Provides a fluent interface for document construction with
    ABNT-compliant formatting.
    """
    
    def __init__(self):
        """Initialize the document builder."""
        self._doc = Document()
        self._configure_page()
    
    def _configure_page(self) -> None:
        """Configure page layout according to ABNT standards."""
        section = self._doc.sections[0]
        section.page_height = Inches(11.69)  # A4: 297mm
        section.page_width = Inches(8.27)    # A4: 210mm
        section.left_margin = Inches(1.18)   # 30mm
        section.right_margin = Inches(0.79)  # 20mm
        section.top_margin = Inches(1.18)    # 30mm
        section.bottom_margin = Inches(1.18) # 30mm
    
    def add_title(self, text: str) -> "DocxBuilder":
        """
        Add a centered title to the document.
        
        Args:
            text: Title text
            
        Returns:
            Self for method chaining
        """
        heading = self._doc.add_heading(text, 0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return self
    
    def add_date(self, date: Optional[datetime] = None) -> "DocxBuilder":
        """
        Add a right-aligned date to the document.
        
        Args:
            date: Date to display. Uses current date if not provided.
            
        Returns:
            Self for method chaining
        """
        date = date or datetime.now()
        para = self._doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = para.add_run(f"Data: {date.strftime('%d/%m/%Y')}")
        run.bold = True
        return self
    
    def add_section(self, number: str, title: str, level: int = 1) -> "DocxBuilder":
        """
        Add a numbered section heading.
        
        Args:
            number: Section number (e.g., "1", "2.1")
            title: Section title
            level: Heading level (1-3)
            
        Returns:
            Self for method chaining
        """
        self._doc.add_heading(f"{number}. {title}", level=level)
        return self
    
    def add_paragraph(
        self,
        text: str,
        bold_prefix: Optional[str] = None,
        italic: bool = False
    ) -> "DocxBuilder":
        """
        Add a paragraph to the document.
        
        Args:
            text: Paragraph text
            bold_prefix: Optional bold prefix text
            italic: Whether to italicize the main text
            
        Returns:
            Self for method chaining
        """
        para = self._doc.add_paragraph()
        
        if bold_prefix:
            para.add_run(bold_prefix).bold = True
        
        run = para.add_run(text)
        run.italic = italic
        
        return self
    
    def add_bullet_list(self, items: List[str], italic: bool = False) -> "DocxBuilder":
        """
        Add a bulleted list to the document.
        
        Args:
            items: List of items to include
            italic: Whether to italicize items
            
        Returns:
            Self for method chaining
        """
        para = self._doc.add_paragraph()
        for i, item in enumerate(items):
            run = para.add_run(f"• {item}")
            run.italic = italic
            if i < len(items) - 1:
                para.add_run("\n")
        return self
    
    def add_table(
        self,
        headers: List[str],
        rows: List[List[str]],
        style: str = "Light Grid Accent 1"
    ) -> "DocxBuilder":
        """
        Add a table to the document.
        
        Args:
            headers: Column headers
            rows: List of row data
            style: Table style name
            
        Returns:
            Self for method chaining
        """
        table = self._doc.add_table(rows=1, cols=len(headers))
        table.style = style
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Add headers
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
        
        # Add data rows
        for row_data in rows:
            row_cells = table.add_row().cells
            for i, cell_value in enumerate(row_data):
                row_cells[i].text = str(cell_value)
        
        self._doc.add_paragraph()  # Space after table
        return self
    
    def add_ranking_table(
        self,
        title: str,
        data: List[Tuple[str, float]],
        value_label: str = "Valor (min)",
        meta: Optional[str] = None,
        description: Optional[str] = None
    ) -> "DocxBuilder":
        """
        Add a ranking table with position, team, and value columns.
        
        Args:
            title: Table title
            data: List of (team, value) tuples
            value_label: Label for the value column
            meta: Optional target/meta description
            description: Optional description text
            
        Returns:
            Self for method chaining
        """
        self._doc.add_heading(title, level=3)
        
        if meta:
            para = self._doc.add_paragraph()
            para.add_run(f"Meta: {meta}").bold = True
        
        if description:
            self._doc.add_paragraph(description)
        
        # Build table rows
        rows = [
            [str(idx + 1), team, f"{value:.2f}"]
            for idx, (team, value) in enumerate(data)
        ]
        
        return self.add_table(
            headers=["Posição", "Equipe", value_label],
            rows=rows
        )
    
    def add_page_break(self) -> "DocxBuilder":
        """
        Add a page break to the document.
        
        Returns:
            Self for method chaining
        """
        self._doc.add_page_break()
        return self
    
    def add_space(self) -> "DocxBuilder":
        """
        Add vertical space (empty paragraph).
        
        Returns:
            Self for method chaining
        """
        self._doc.add_paragraph()
        return self
    
    def save(self, path: str) -> None:
        """
        Save the document to a file.
        
        Args:
            path: File path to save to
        """
        self._doc.save(path)
        logger.info(f"Document saved to: {path}")
    
    @property
    def document(self) -> Document:
        """Get the underlying Document object."""
        return self._doc
